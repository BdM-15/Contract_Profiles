# Description: The OSBP module contains resusable and commonly used fucntions, classes, dictionaries...etc.  These are used to provide insights into OSBP-RI data anlaytics to include predictive analysis and machine learning and generative AI.  

# Import the most commonly used libraries to support the commonly used functions, classes, dictionaries, etc. 
import pandas as pd
import numpy as np 
import datetime
import os 
import glob 
import openpyxl
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.dml import MSO_THEME_COLOR_INDEX



common_folders = {
    #This is the initial data pull from the data sources. This is the first step in the data pipeline.  It will contain both the acc_ri and army data sources. POTENTIALLY MAKE THIS ITS OWN SCRIPT TO BE RUN MONTHLY SINCE READING THE ARMY DATA SOURCE TAKES ABOUT 30 MINUTES.  THIS WILL ALLOW ONLY ONE SPREADSHEET TO BE NEEDED.
    'raw_data_folder': r'C:\GitHub\contract_profiles\data\raw',
    
    'raw_acc_ri_data_source_file': r'C:\GitHub\contract_profiles\data\raw\raw_acc_ri_data_source_file.xlsx',
    'raw_army_data_source_file': r'C:\GitHub\contract_profiles\data\raw\raw_army_data_source_file.xlsx',
    
    # The interime folder is where the clean and transformed data source files is stored and used to build all insights target lists from.  This is the most important folder in the data pipeline.
    'interim_data_source_folder': r'C:\GitHub\contract_profiles\data\interim',
    'interim_data_source_file': r'C:\GitHub\contract_profiles\data\interim\acc-ri_data_source.csv',
    'interim_army_data_source_folder': r'C:\GitHub\contract_profiles\data\interim',
    'interim_army_data_source_file': r'C:\GitHub\contract_profiles\data\interim\army_data_source.csv',
    
    # The processed folder is the location where all the target lists are created and stored.  These target lists is what the contract profiles are built from.
    'inights_folder': r'C:\GitHub\contract_profiles\data\processed',
    
    'insight_unrestricted_awarded_to_sb_folder': r'C:\GitHub\contract_profiles\data\processed\unrestricted_awarded_to_sb',
    'insight_unrestricted_awarded_to_sb_file': r'C:\GitHub\contract_profiles\data\processed\unrestricted_awarded_to_sb\insights_target1.csv',
    
    'insight_sbsa_with_potential_for_socio_set_aside_folder': r'C:\GitHub\contract_profiles\data\processed\sbsa_with_potential_for_socio_set_asides',
    'insight_sbsa_with_potential_for_socio_set_aside_file': r'C:\GitHub\contract_profiles\data\processed\sbsa_with_potential_for_socio_set_asides.csv',
    
    'insight_8a_with_exit_before_expiration_folder': r'C:\GitHub\contract_profiles\data\processed\8a_with_exit_before_expiration',
    'insight_8a_with_exit_before_expiration_file': r'C:\GitHub\contract_profiles\data\processed\8a_with_exit_before_expira.csv',
    
    # The reports folder is where the initial contract profiles are stored.  This is the final step in the data pipeline before using them to support proactive market research.
    'contract_profiles_folder': r'C:\GitHub\contract_profiles\reports\contract_profiles',
    'contract_profiles_template_file': r'C:\GitHub\contract_profiles\reports\contract_profiles\template_contract_profile.docx',
    'completed_profiles_folder': r'C:\GitHub\contract_profiles\reports\contract_profiles\completed_profiles',
    
    # The references folder is where the reference data is stores.  This data is used to support the SB Profile Analysis in the contract profiles.
    # The Underepresented NAICS list comes from the SBA website and is used to identify NAICS codes that are underrepresented by WOSBs and EDWOSBs.  This list is used to identify contracts that are set aside for WOSBs and EDWOSBs.
    'wosb_naics_list' : r'C:\GitHub\contract_profiles\references\wosb_naics\wosb_naics_list.csv',
    # The NMR Waiver list comes from the SBA website and is used to identify NAICS codes that have a waiver for the Non-Manufacturer Rule.  This list is used to identify contracts that have a waiver for the NMR.
    'nmr_waiver_list' : r'C:\GitHub\contract_profiles\references\nmr_class_waivers\nmr_waiver_list.csv',
    # The Size Standard list comes from the SBA website and is used to identify the size standard for each NAICS code.
    'size_standard_list' : r'C:\GitHub\contract_profiles\references\size_standards\size_standards_list.csv',
    'forecast_folder' : r'C:\GitHub\contract_profiles\references\forecast_listing',
    'osbp_forecast_file' : r'C:\GitHub\contract_profiles\references\forecast_listing\osbp_forecast.csv',
    'amc_forecast_file' : r'C:\GitHub\contract_profiles\references\forecast_listing\amc_forecast.csv',
    'hyperlinks_folder' : r'C:\GitHub\contract_profiles\references\hyperlinks',
    'hyperlinks_file' : r'C:\GitHub\contract_profiles\references\hyperlinks\hyperlinks_listing.csv',
}
# This function finds the latest file in a folder based on the file pattern.  It will determine if the file is .xlsx or .csv and read the file accordingly.
def find_latest_file(folder_path: str, file_pattern: str) -> str:
    """
    Find the latest file in a folder based on the file pattern.

    Args:
    folder_path (str): The path to the folder where the files are located.
    file_pattern (str): The pattern to search for the latest file.

    Returns:
    str: The path to the latest file found in the folder.
    """
   
    file_list = glob.glob(os.path.join(folder_path, file_pattern))

    if file_list:
        latest_file = max(file_list, key=os.path.getctime)
        # Print the latest file to the console for confirmation
        # print(f"Latest file: {latest_file}") 
        return latest_file # Return the path to the latest file
    else:
        raise FileNotFoundError(f"No file matching the pattern {file_pattern} found in {folder_path}.")
    
#This function will identify any older files than the most recent file in the raw data folder and move them to the archive folder.  If an archive folder does not exist, it will create one.
def archive_folder_files(current_folder: str, file_list: list, latest_file: str) -> None:
    """
    Move older files in the raw data folder to the archive folder.

    Args:
    current_folder (str): The path to the data folder we want to have an archive.
    file_list (list): A list of files in the data folder to determine if any older files exist.
    latest_file (str): The path to the latest file in the data folder.

    Returns:
    None
    """
    archive_folder = os.path.join(current_folder, 'archive') # Define the archive folder location

    #  Create the archive folder if it does not exist
    if not os.path.exists(archive_folder):
        os.makedirs(archive_folder)
    for file in file_list: # Loop through the files in the raw data folder
        if file != latest_file: # Check if the file is not the latest file
            os.rename(file, os.path.join(archive_folder, os.path.basename(file))) # Move the older files to the archive folder

def remove_old_files(folder_path: str, file_pattern: str) -> None:
    """
    Check for the latest file in the folder and remove older files in a folder based on the file pattern.

    Args:
    folder_path (str): The path to the folder where the files are located.
    file_pattern (str): The pattern to search for the files to remove.

    Returns:
    None
    """
    # Find the latest file in the folder based on the file pattern and define it as the latest_file_to_keep
    latest_file_to_keep = find_latest_file(folder_path, file_pattern)
    
    # Check for any folders older than the latest file and remove them
    for file in glob.glob(os.path.join(folder_path, file_pattern)):
        if file != latest_file_to_keep:
            os.remove(file)

# This function will find the latest file and convert the raw data file from .xlsx to .csv and move it to the interim data folder.  It will also log the conversion details in a log file.
def convert_raw_data_from_excel_to_csv() -> str:
    """
    Convert the raw data file from .xlsx to .csv and move it to the interim data folder.
    https://x.com/i/grok/share/nUScJIxpbmfLVBcnSCQqO8xQ1
    
    Args:
    None

    Returns:
    str: The path to the converted file in the interim data folder.
    """
    # Define the raw data folder location and the file pattern to search for the latest file in the raw data folder.
    raw_data_folder = 'raw_data_folder'
    
    # Define the destination folder and the naming convention of the converted file
    interim_data_folder = 'interim_data_folder'
     
    # Get today's date
    today_date = datetime.datetime.now().strftime('%Y-%m-%d') 
    
    # Define the file pattern to search for the latest file in the raw data folder and create a list of all the xlsx files in the raw data folder.
    raw_file_pattern = os.path.join(raw_data_folder, '*.xlsx') 
    raw_file_list = glob.glob(raw_file_pattern)
    
    # Find the latest file in the raw data folder based on the file pattern
    latest_file = find_latest_file(raw_data_folder, '*.xlsx')

    # Call the archive_folder_files function if more than one file exists in the folder to move older files to the archive folder
    if len(raw_file_list) > 1: 
        archive_folder_files(raw_data_folder, raw_file_list, latest_file)

    # Define the destination folder and the naming convention of the converted file
    interim_data_file = os.path.join(interim_data_folder, f"data_source.csv")
    
    # Read the latest raw data file 
    df = pd.read_excel(latest_file)
    
    # Convert the latest raw data file to a CSV file and save it to the cleansed data folder
    df.to_csv(interim_data_file, index=False)
    
    
    # Print the conversion details to the console
    print(f"{latest_file} was successfully converted to {interim_data_file} on {today_date}.")

    # Log the conversion details
    log_file = os.path.join(raw_data_folder, 'OSBP-RI_data_conversion_log.txt') 
    with open(log_file, 'a') as log:
        log.write(f"The raw data to for the OSBP Insights was converted on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\nRaw data file: {latest_file}\nConverted file: {interim_data_file}\nConverted csv file location: {interim_data_folder}\n\n")
    print(f"Conversion details logged in {log_file}")
    
    return interim_data_file


# This functions provides a general cleanse of the data.  It removes any rows and columns that are completely empty, removes any duplicate rows, reduces the memory usage of the DataFrame by converting columns with object dtype to category dtype, removes unique_values "Modification", "MATOC", "SATOC" from the 'Contract Action Type' column, replaces any blank values in "10N Type Set Aside Description" with "NO SET ASIDE USED.", converts the 'Current Completion' column to datetime to match format with today_date, calculates the number of months remaining to complete the contract and inputs the number of months remaining in a new 'Months Remaining' column.
def clean_and_transform_data_for_contract_profiles(df, destination_folder: str):
    """
    Clean and transforme the data. Ensure the proper data types are used; rename any columns to align with desired result.  Also, save any formatting for the point in which it will be analyzed.  For example, don't add currency formatting at this point.  Just make sure numbers are numeric, dates are dates, and text is text.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be cleaned.

    Returns:
    pd.DataFrame: The cleaned DataFrame.
    """
  
    # Read the latest CSV file
    df = pd.read_csv(df)
    
    # Remove any rows and columns that are completely empty
    df = df.dropna(how='all')
    df = df.dropna(axis=1, how='all')

    # Remove any duplicate rows
    df = df.drop_duplicates()

    # Reduce the memory usage of the DataFrame by converting columns with object dtype to category dtype
    df = df.astype({col: 'category' for col in df.select_dtypes('object').columns})
    
    # Rename column "13GG Legal Business Name (UEI)" to "Awardee"
    df = df.rename(columns={"13GG Legal Business Name (UEI)": "Awardee"})
    
    # Rename column "10N Type Set Aside Description" to "Type Set Aside Description"
    df = df.rename(columns={"10N Type Set Aside Description": "Type Set Aside Description"})
    
    # Remove the . from the values in the 'Type Set Aside Description' column
    df['Type Set Aside Description'] = df['Type Set Aside Description'].str.replace('.', '')
    
    # Rename column "6M Desription of Requirement" to "Requirements Description"
    df = df.rename(columns={"6M Description of Requirement": "Requirements Description"})
        
    # Rename column "Current Completion Date" to "Exipriation"
    df = df.rename(columns={"Current Completion Date": "Expiration"})
    
    # Rename column "Small Business Actions" to "Size Status"
    df = df.rename(columns={"Small Business Actions": "Size Status"})
    
    # Within Size Status column, replace "0" with "OTSB" and "1" with "SB"
    df['Size Status'] = df['Size Status'].replace({0: 'OTSB', 1: 'SB'})
    
    # Rename column "Small Business Dollars" to "SB Dollars"
    df = df.rename(columns={"Small Business Dollars": "SB Dollars"})
        
    # Remove any value after the six digit in the NAICS column
    df['NAICS'] = df['NAICS'].astype(str).str[:6]   
    
    # # Remove unique_values "Modification", "MATOC", "SATOC" from the 'Contract Action Type' column
    # df = df[~df['Contract Action Type'].str.upper().isin(["MODIFICATION", "MATOC", "SATOC"])]

    # If a blank value exsits in "10N Type Set Aside Description" replace it with "NO SET ASIDE USED."
    df['Type Set Aside Description'] = df['Type Set Aside Description'].fillna("NO SET ASIDE USED")

    # Convert the 'Expiration' column to datetime to match format with today_date and make it only 10 characters long
    df['Expiration'] = df['Expiration'].str[:10]
    df['Expiration'] = pd.to_datetime(df['Expiration'], errors='coerce')
    
    # Convert the "Award Date" column to datetime to match format with today_date
    df['Award Date'] = pd.to_datetime(df['Award Date'], errors='coerce')

    # For each row, calulate the number of months remaining to complete the contract and input the number of months remaining in the 'Months Remaining' column
    df['Months Remaining'] = (df['Expiration'] - pd.Timestamp.today()).dt.days // 30

    # Add a new column next to the 'Current Completion' column and insert the 'Months Remaining' column
    df.insert(df.columns.get_loc('Expiration') + 1, 'Months Remaining', df.pop('Months Remaining'))
    
    # Save the cleaned data copy to the cleansed data folder using the destination_folder argument
    df.to_csv(os.path.join(destination_folder, 'data_source.csv'), index=False)

    # If any older files exist call the remove_old_files function and remove them
    remove_old_files(destination_folder, '*.csv')
    
    return df
    
    # Example usage of the clean_data function
    # df = clean_data(df, 'path/to/destination_folder')
    
# This function will make a copy of the cleaned data and move it to the "targets" folder which is in the "processed" data folder.  This data will be used to sort for the targeted contracts within the next 6-18 months for contract profile development.  This includes Full and Opens that were awarded to SBs, SBSAs with potential for socio set asides.
def insights_unrestricted_awarded_to_sb() -> None:
    """
    Process the data to sort for the targeted contracts.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.

    Returns:
    None
    """
    # Define the targets folder location. If the targets folder does not exist, create it.
    if not os.path.exists(common_folders['insights_target1_folder']):
        os.makedirs(common_folders['insights_target1_folder'])
    
    # Define the cleansed data source file to create a DataFrame (df)
    df = pd.read_csv(find_latest_file(common_folders['cleansed_data_source_folder'], '*.csv'))
    
    # Filter the DataFrame for rows where 'Contract Action Type' is not "MODIFICATION".
    df = df[df['Contract Action Type'].str.upper() != "MODIFICATION"]
           
    # # Keep "ACC-RI" in the "organization" column.  Remove any other values.
    # df = df[df['Organization'] == 'ACC-RI']
    
    # Ensure the NAICS code is only first six digits
    df['NAICS'] = df['NAICS'].astype(str).str[:6]
    
    # Convert the "SB Dollars" column to currency format
    df['SB Dollars'] = df['SB Dollars'].map('${:,.2f}'.format)
    
     # Within Size Status column, replace "0" with "OTSB" and "1" with "SB"
    df['Size Status'] = df['Size Status'].replace({0: 'OTSB', 1: 'SB'})
   
    # Remove any rows that are not "NO SET ASIDE USED" or blank in the "10N Type Set Aside Description" column
    df = df[df['Type Set Aside Description'].isin(["NO SET ASIDE USED", ""])]
    
    # Remove any rows with Months reaminining less than 6 months and more than 18 months
    df = df[(df['Months Remaining'] >= 6) & (df['Months Remaining'] <= 18)]
    
    # Remove any rows with "OTSB in the "Size Status"
    df = df[df['Size Status'] != "OTSB"]
    
    # Sort "Months Remaining" in ascending order
    df = df.sort_values(by='Months Remaining', ascending=True)
    
    # Save the target processed data copy to the targets data folder
    df.to_csv(os.path.join(common_folders["insights_target1_folder"], 'insights_target1.csv'), index=False)
    
    # If any older files exist call the remove_old_files function and remove them
    remove_old_files(common_folders['insights_target1_folder'], '*.csv')
    
# Create a dictionary of Contract Profile Data Elements to populate the Contract Profile Data Elements table in the Contract Profile Word Document.
contract_profile_data_elements = {
    "1A" : "Contract No",
    "1A2" : "Order No",
    "1B" : "Modification No",
    "2A" : "Award Date",
    "2B" : "Effective Date",
    "2D" : "Expiration",
    "6A" : "Contract Type",
    "6M" : "Requirements Description",
    "8G" : "NAICS",
    "8N" : "Bundling",
    "9G" : "Place of Performance",
    "10C" : "Limited Competition",
    "10D" : "Number of Offerors",
    "10N" : "Type Set Aside Description",
    "11B" : "SubK Plan",
    "12A" : "GWAC", #Includes Agency IDIQs
    "13GG" : "Awardee",
    "CAT" : "Contract Action Type",
    "CPAR" : "CPARS Rating",
    "ESRS" : "SubK Achievement",
    "SB$" : "SB Dollars",
    "STS" : "Size Status", #originally "Small Business Actions" (0 = OTSB and 1 = SB)
    "FC" : "Forecast No", # Identify the forecast solicitation/PANCOC number
    "PCF" : "PCF Cabinet", #Provide link to PCF Cabinet
    # "FCL" : "Forecast Link", #Provide link to Forecast PCF Cabinet
}
   
sb_categories = {   
    "sb" : "SB",
    "d"  : "SDB",
    "w"  : "WOSB",
    "ew" : "EDWOSB",
    "v" : "VOSB",
    "sv" : "SDVOSB",
    "8a" : "8(a)",
    "hz" : "HUBZone",
}

sb_profile_analysis_elements = {
    "IT" : "IT Buy", 
    "ITSS" : "IT Services SONA",
    "SComp" : "Strong Competition",
    "SStd" : "Size Standard",
    "TNAICS:" : "Top NAICS",
    "TgtNAICS" : "Target NAICS",
    "WElg" : "WOSB Eligible",
    # "SoSSElg" : "Socio SS Eligible",
    "SNAICS" : "Strong NAICS",
    "WNAICS" : "Weak NAICS",
    "RIAwd" : "ACC RI Awards", #Awards that went to small under the identified NAICS
    "AAwd" : "All ACC Awards", #All awards made by ACC across the enterprise
    "SkR" : "Subcontract MQRs Realistic",
    "AwdSB" : "Awardee SB",
    "AwdSoc" : "Awardee Socio",
    "Mult" : "Multiple Products or Services",
    "NMRW" : "NMR Waiver Available", #Does an NMR waiver exist based on NAICS
    "NMRP" : "NMR Potential", #Potential for NMR based on requirements
    "FinC" : "Financial Risk", #Financial risk to industry based on distribution of SB awards under identified NAICS
    # "Rmks" : "Remarks"
}

# def add_hyperlink(paragraph, text, url):
#     """
#     Add a hyperlink to a paragraph.

#     Args:
#     paragraph (docx.text.paragraph.Paragraph): The paragraph to add the hyperlink to.
#     text (str): The text to display for the hyperlink.
#     url (str): The URL to link to.

#     Returns:
#     None
#     """

#     # Create a new run for the hyperlink
#     run = paragraph.add_run()
#     run.text = text
    
#     # Create the hyperlink element
#     hyperlink = run._r
#     hyperlink.add_rPr()
#     hyperlink.rPr.add_hlinkClick(url, qn('r:id'), qn('r:uri'))
#     run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
#     run.font.underline = True

#     # Assuming you want to put the hyperlink in the first cell of the first row
#     cell = table.cell(0, 0)
#     hyperlink_text = "Click here"  # The string you want to display
#     hyperlink_url = "http://www.example.com"  # The URL

#     # Add the hyperlink to the paragraph inside the cell
#     add_hyperlink(cell.paragraphs[0], hyperlink_text, hyperlink_url)

# This helper function to change the font size of the non-header cells to size 9 Calibri.
def set_cell_font(cell, font_name='Calibri', font_size=9, bold=False, color=None):
    """
    Set the font for a table cell.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color) 

def check_size_standard(df, contract_no) -> str:
    """
    Check if the NAICS code value is present in the Size Standard listing (size_standard_list.xlsx).

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: Value from from the 'Size standards in millions of dollars' column.  If 'Size standards in millions of dollars' is not present, return value from 'Size standards in number of employees' column.  If the NAICS value is not present, return "No".
    """
    
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]

    # Read the Size Standard listing
    size_standard_df = pd.read_csv(common_folders['size_standard_list'])
    
    # Remove any value after the six digit in the NAICS Code column
    size_standard_df['NAICS Code'] = size_standard_df['NAICS Codes'].astype(str).str[:6]
    
    # Make sure the "Size standards in millions of dollars" column is string type
    size_standard_df['Size standards in millions of dollars'] = size_standard_df['Size standards in millions of dollars'].astype(str)
    
    # Make sure the "Size standards in number of employees" column is string type
    size_standard_df['Size standards in number of employees'] = size_standard_df['Size standards in number of employees'].astype(str)
    
    # Check if naics is in the size_standard_df['NAICS Codes'] column. Search the 'Size standards in millions of dollars' column for a value. If a value exists return that value. If no value found move over one column and return its value.  Otherwise, return naics not found message
    if naics in size_standard_df['NAICS Codes'].values:
        if size_standard_df.loc[size_standard_df['NAICS Code'] == naics, 'Size standards in millions of dollars'].values[0] != 'nan':
            return str(size_standard_df.loc[size_standard_df['NAICS Code'] == naics, 'Size standards in millions of dollars'].values[0]).strip() + "M"
        else:
            return str(size_standard_df.loc[size_standard_df['NAICS Code'] == naics, 'Size standards in number of employees'].values[0]).strip() + " Employees"
    else:
        return f'{naics} not found'  
     
def check_wosb_naics(df, contract_no) -> str:
    """
    Check if the NAICS code value is present in the Underrepresented WOSB NAICS listing (wosb_naics_list.xlsx).

    Args
    Datframe (df): The NAICS value to be checked based on the current contract number being processed.  Should be insight_target.csv.
    
    Contract_no (str): The contract number being processed in create_contract_profiles() function.

    Returns:
    str: "WOSB" or "EDWOSB" from the 'Set-Aside' column if the NAICS value is present, "No" otherwise.
    """
    
    # Select the NAICS value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
        
    # Read the WOSB NAICS listing
    wosb_naics_df = pd.read_csv(common_folders['wosb_naics_list'])
    
    # Remove any value after the six digit in the NAICS Code column
    wosb_naics_df['NAICS Code'] = wosb_naics_df['NAICS Code'].astype(str).str[:6]
    
    # Make sure the "Set-aside" column is string type
    wosb_naics_df['Set-aside'] = wosb_naics_df['Set-aside'].astype(str)
    
    # Check if naics is in the wosb_naics_df['NAICS Code'] column. If yes, return the value in the 'Set-aside' column. If no, return "No"
    if naics in wosb_naics_df['NAICS Code'].values:
        return wosb_naics_df.loc[wosb_naics_df['NAICS Code'] == naics, 'Set-aside'].values[0]
    else:
        return "No"

def check_if_awardee_sb(df, contract_no) -> str:
    """
    Check if the awardee is a small business based on the 'Size Status' column.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: "SB" if the awardee is a small business, "No" otherwise.
    """
    
    # Select the 'Size Status' value from the DataFrame based on the current contract number being processed
    size_status = df.loc[df['Contract No'] == contract_no, 'Size Status'].values[0]
    
    # Check if the 'Size Status' value is "SB". If yes, return "Yes". If no, return "No"
    if size_status == "SB":
        return "Yes"
    else:
        return "No"
    
def check_awardee_socioeconomic_status(df, contract_no) -> str:
    """
    Check if the awardee is a socio-economic status based on the 'Size Status' column.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: All socio categories they are identified as based on the "WOSB", "EDWOSB", "VOSB", "SDVOSB", "8(a)", "HUBZone", or "No" based on the socio-economic status.  If multiple categories are identified, they will be separated by a comma.
    """
    
    # Select the 'Size Status' value from the DataFrame based on the current contract number being processed
    socioeconomic_status = df.loc[df['Contract No'] == contract_no, 'Size Status'].values[0]
    
    # Create a list of all socioeconomic status the Awardee is.  Check the "SDB Concern Action", "Service Disabled Veterans Actions", "Women Owned Actions", and "HUB Zone Actions" columns to determine size status.  If the value is "1" then the Awardee is that socio-economic status and add it to the list. After reviewing all columns, return the list of socio-economic status as string seperated by a comma based on list and if list is empty return "None".
    socioeconomic_categories = []
    
    # Define a dictionary mapping column names to their corresponding categories
    socioeconomic_columns = {
        'SDB Concern Actions': 'SDB',
        'Service Disabled Veterans Actions': 'SDVOSB',
        'Women Owned Actions': 'WOSB',
        'HUB Zone Actions': 'HUBZone'
    }

    # Iterate through the dictionary and check each column
    for column, category in socioeconomic_columns.items():
        if df.loc[df['Contract No'] == contract_no, column].values[0] == 1:
            socioeconomic_categories.append(category)
    
    if socioeconomic_categories:
        return ', '.join(socioeconomic_categories)
    else:
        return "None"
 
def check_if_nmr_waiver_available(df, contract_no) -> str:
    """
    Check if an NMR waiver exists based on the NAICS code from the current contract being processed.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: "Yes" if an NMR waiver exists, "No" otherwise.
    """
    
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
        
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
    
    # Define path for nmr_waiver_list
    # wosb_naics_path = common_folders['wosb_naics_list']
    
    # Read the nmr_waiver_list (needs to be saved as CSV UTF-8)
    nmr_waiver_list_df = pd.read_csv(common_folders['nmr_waiver_list'])
    
    # Remove any value after the six digit in the NAICS Code column
    nmr_waiver_list_df['NAICS CODE'] = nmr_waiver_list_df['NAICS CODE'].astype(str).str[:6]
    
    # Make sure the "NAICS DESCRIPTOR" column is string type
    nmr_waiver_list_df['NAICS DESCRIPTOR'] = nmr_waiver_list_df['NAICS DESCRIPTOR'].astype(str)
    
    # Check if naics is in the nmr_waiver_list_df['NAICS Code'] column. If yes, return the value in the 'Set-aside' column. If no, return "No"
    if naics in nmr_waiver_list_df['NAICS CODE'].values:
        return "Yes"
        # return nmr_waiver_list_df.loc[nmr_waiver_list_df['NAICS CODE'] == naics, 'NAICS DESCRIPTOR'].values[0]
    else:
        return "No"
    
def check_acc_ri_awards(df, contract_no) -> str:
    """
    Get the number of awards made by ACC-RI to small businesses based on the NAICS code from the current contract being processed.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    int: The number of awards made.
    """
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
    # print(naics)
    # print("")
    
    # Read into df the data source file that will be used to develop the percentiles from the SB Dollars column 
    acc_ri_awards_df = pd.read_csv(common_folders['cleansed_data_source_file'])
    
    # Filter df to only show Contract Action Types that are not "MODIFICATION", "SATOC", and "MATOC"
    acc_ri_awards_df = acc_ri_awards_df[~acc_ri_awards_df['Contract Action Type'].str.upper().isin(["MODIFICATION", "MATOC", "SATOC"])]
    # print("Dataframe before filter:", acc_ri_awards_df)
    # print("")
    
    # # Ensure the NAICS column is of the same type as naics
    # acc_ri_awards_df['NAICS'] = acc_ri_awards_df['NAICS'].astype(type(naics))
   
    # Ensure the NAICS column is of the same type as naics and only six digits are used
    acc_ri_awards_df['NAICS'] = acc_ri_awards_df['NAICS'].astype(str).str[:6].astype(type(naics))

    # Ensure the Size Status column is of the same type as "SB"
    acc_ri_awards_df['Size Status'] = acc_ri_awards_df['Size Status'].astype(str).str.strip()

    acc_ri_awards_df = acc_ri_awards_df.loc[(acc_ri_awards_df['NAICS'] == naics) & (acc_ri_awards_df['Size Status'] == "SB")]
    # print("Dataframe after filter:", acc_ri_awards_df)
    # print("")
    
    # If dataframe is empty return "0", else return the number of rows remaining after filter to get the award count
    if acc_ri_awards_df.empty:
        return str(0)
    else:
        return str(acc_ri_awards_df.shape[0])
    
def check_all_acc_awards(df, contract_no):
    """
    Get the number of awards made by the Army enterprise to small businesses based on the NAICS code from the current contract being processed.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    int: The number of awards made.
    """
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
    # print(contract_no)
    # print(naics)
    # print("")
    
    # Only read the 'NAICS', 'Size Status', and ' columns into df the data source file that will be used to determine award count.
    acc_awards_df = pd.read_csv(common_folders['cleansed_all_army_data_source_file'], usecols=['NAICS', 'Size Status', 'Contract Action Type'])
    # acc_awards_df = pd.read_csv(common_folders['cleansed_all_army_data_source_file'])
  
    # Remove unique_values "Modification", "MATOC", "SATOC" from the 'Contract Action Type' column
    acc_awards_df = acc_awards_df[~acc_awards_df['Contract Action Type'].str.upper().isin(["MODIFICATION", "MATOC", "SATOC"])]
    
    # # Rename column "Small Business Actions" to "Size Status"
    # acc_awards_df = acc_awards_df.rename(columns={"Small Business Actions": "Size Status"})
    
    # # Within Size Status column, replace "0" with "OTSB" and ">0" with "SB"
    # acc_awards_df['Size Status'] = acc_awards_df['Size Status'].replace({0: 'OTSB', 1: 'SB'})
                      
    # Ensure the NAICS column is of the same type as naics and only six digits are used
    acc_awards_df['NAICS'] = acc_awards_df['NAICS'].astype(str).str[:6].astype(type(naics))
    # acc_awards_df['NAICS'] = acc_awards_df['NAICS'].astype(type(naics))
        
    # Ensure the Size Status column is of the same type as "SB" and only six digits are used
    acc_awards_df['Size Status'] = acc_awards_df['Size Status'].astype(str).str.strip()
    
    # print("Dataframe before filter:", acc_awards_df)
    # print("")
    # Filter the acc_awards_df based on the naics identified from the contract being processed
    acc_awards_df = acc_awards_df.loc[(acc_awards_df['NAICS'] == naics) & (acc_awards_df['Size Status'] == "SB")]
    # print("Dataframe after filter:", acc_awards_df)
    # print("")
    
    # If dataframe is empty return "0", else return the number of rows remaining after filter to get the award count
    if acc_awards_df.empty:
        return str(0)
    else:
        return str(acc_awards_df.shape[0])
    
def check_financial_risk(df, contract_no) -> str:
    """
    Check the financial risk to industry based on the distribution of SB dollars against the identified NAICS.
    Percentiles will be 50% for Low, 75% for Medium, and 90% and above for High. 

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: "High", "Medium", "Low", or "No" based on the financial risk.
    """
    
    # Read into df the data source file that will be used to develop the percentiles from the SB Dollars column 
    # Use the below for testing the logic.  It's a smaller data sample.
    # cleansed_file_df = pd.read_csv(common_folders['cleansed_data_source_file'])
    cleansed_file_df = pd.read_csv(common_folders['cleansed_all_army_data_source_file'], usecols=['NAICS', 'Size Status', 'SB Dollars'])
    
    # Define the NAICS values from the DataFrame argument (NOT THE FULL RAW CLEANSED FILE) based on the current contract number being processed.  This will be used to sort the cleansed_file_df to get the percentiles.
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # # Print the unique values in the 'NAICS' column to verify the expected value is present
    # print("Unique NAICS values in the DataFrame:", df['NAICS'].unique())
    # print("")
    # print(contract_no)
    # print(naics)
    # print("")

    # Make sure the 'SB Dollars' column is numeric from the DataFrame argument (NOT THE FULL RAW CLEANSED FILE).  This value will be used to compare against the percentiles from the cleansed_file_df.
    # Ensure 'SB Dollars' column is string before removing currency formatting
    df['SB Dollars'] = df['SB Dollars'].astype(str)

    # Remove currency formatting from the SB Dollars column
    df['SB Dollars'] = df['SB Dollars'].str.replace('$', '').str.replace(',', '')

    # Convert the SB Dollars column to numeric
    df['SB Dollars'] = pd.to_numeric(df['SB Dollars'], errors='coerce')

    # Make sure the 'SB Dollars' column is numeric from the DataFrame argument (NOT THE FULL RAW CLEANSED FILE). This value will be used to compare against the percentiles from the cleansed_file_df.
    sb_dollars = df.loc[df['Contract No'] == contract_no, 'SB Dollars'].values[0]
    # print('SB Dollars before conversion: ', sb_dollars)
    
    # sb_dollars = pd.to_numeric(df.loc[df['Contract No'] == contract_no, 'SB Dollars'].values[0], errors='coerce')
    # print('SB Dollars after conversion: ', sb_dollars)
    # print("")

    # Filter the cleased_file_df based on the naics identified from the contract being processed and print the filtered result to check if it matches the expected rows
    filtered_df = cleansed_file_df.loc[(cleansed_file_df['NAICS'] == naics) & (cleansed_file_df['Size Status'] == "SB")]
    # print("Filtered DataFrame based on NAICS:", filtered_df)

    # Check if the 'SB Dollars' column contains the expected values from the full raw cleansed file (latest_cleansed_file_df)
    if not filtered_df.empty:
        # Print the 'SB Dollars' column before conversion
        # print("SB Dollars column before conversion:", filtered_df['SB Dollars'])
        
        # Convert 'SB Dollars' column to numeric values
        filtered_df.loc[:, 'SB Dollars'] = pd.to_numeric(filtered_df['SB Dollars'], errors='coerce')
        
        # Print the 'SB Dollars' column after conversion
        # print("SB Dollars column after conversion:", filtered_df['SB Dollars'])
        
        sb_dollars_values = filtered_df['SB Dollars'].values
        # print("SB Dollars Values:", sb_dollars_values)
        
        # Calculate the percentiles
        p50 = np.percentile(sb_dollars_values, 50)
        p75 = np.percentile(sb_dollars_values, 75)
        p90 = np.percentile(sb_dollars_values, 90)
        
        # print(f"50th Percentile (Low Risk): ${p50}")
        # print(f"75th Percentile (Medium Risk): ${p75}")
        # print(f"90th Percentile (High Risk): ${p90}")
        
        # Determine the risk level for the current contract's SB Dollars
        if sb_dollars <= p50:
            risk_level = "Low Risk"
        elif sb_dollars <= p75:
            risk_level = "Medium Risk"
        else:
            risk_level = "High Risk"
        
        return risk_level

def check_targeted_naics(df, contract_no) -> str:
    """
    Check if the NAICS code value is one of the Targeted Econcomic Sector (first two digits of NAICS) identified.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str:  If NAICS is present, return "Yes". Otherwise, return "No".
    """
    
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first two digits are used
    naics = str(naics)[:2]
    
    # Define the targeted NAICS values.  Based on OSBP or Contracting goals.
    targeted_naics = ['33', '51', '54']
    
    # Check if the 'NAICS' value is in the targeted_naics list. If yes, return "Yes". If no, return "No"
    if naics in targeted_naics:
        return "Yes"
    else:
        return "No"

def check_top_naics(df, contract_no) -> str:
    """
    Check if the NAICS code value is one of the Top 25 NAICS identified by the amount of SB Actions or SB Dollars.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str:  If NAICS is present, return "Yes". Otherwise, return "No".
    """
    
    cleansed_file_df = pd.read_csv(common_folders['cleansed_data_source_file'], usecols=['NAICS', 'Size Status', 'SB Dollars'])
    
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
    
  
    # Get the count for top 30% of NAICS based on SB Dollars or SB actions
    top_naics_count = 25
    
    
    # Make sure the 'SB Dollars' column is numeric from the DataFrame argument (NOT THE FULL RAW CLEANSED FILE).  This value will be used to compare against the percentiles from the cleansed_file_df.
    # Ensure 'SB Dollars' column is string before removing currency formatting
    cleansed_file_df['SB Dollars'] = cleansed_file_df['SB Dollars'].astype(str)

    # Remove currency formatting from the SB Dollars column
    cleansed_file_df['SB Dollars'] = cleansed_file_df['SB Dollars'].str.replace('$', '').str.replace(',', '')

    # Convert the SB Dollars column to numeric
    cleansed_file_df['SB Dollars'] = pd.to_numeric(cleansed_file_df['SB Dollars'], errors='coerce')
    
    # Ensure the 'NAICS' column is of the same type as naics and only six digits are used
    cleansed_file_df['NAICS'] = cleansed_file_df['NAICS'].astype(str).str[:6].astype(type(naics))
    
    # Filter the cleased_file_df based on the naics identified from the contract being processed and print the filtered result to check if it matches the expected rows
    # filtered_df = cleansed_file_df.loc[(cleansed_file_df['NAICS'] == naics) & (cleansed_file_df['Size Status'] == "SB")]
    filtered_df = cleansed_file_df.loc[(cleansed_file_df['Size Status'] == "SB")]
    # print("Filtered DataFrame based on NAICS and Size Status:", filtered_df)
    # print("")
    
    # Define the targeted NAICS values based on the top 30% of unique NAICS based on total SB Dollars
    filtered_df = filtered_df.groupby('NAICS').agg({'SB Dollars': 'sum'}).reset_index()
    filtered_df = filtered_df.sort_values(by=['SB Dollars'], ascending=False)
    top_naics_by_dollars = filtered_df['NAICS'].head(top_naics_count).tolist()
    # print("Top NAICS by SB Dollars:", top_naics_by_dollars)
    # print("")
    
    # Define the targeted NAICS values based on the top 30% of unique NAICS based on total SB Actions (determined by "Size Status" column).
    filtered_df = cleansed_file_df.loc[(cleansed_file_df['Size Status'] == "SB")]
    filtered_df = filtered_df.groupby('NAICS').agg({'Size Status': 'count'}).reset_index()
    filtered_df = filtered_df.sort_values(by=['Size Status'], ascending=False)
    top_naics_by_actions = filtered_df['NAICS'].head(top_naics_count).tolist()
    # print("Top NAICS by SB Actions:", top_naics_by_actions)
    # print("")
    
    # top_naics = ['541611', '541512', '541330']
    
    # Check if the 'NAICS' value is in either top_naics_by_actions or top_naics_by_actions. If yes in either one, return "Yes". If not in either, return "No"
    if naics in top_naics_by_dollars or naics in top_naics_by_actions:
        return "Yes"
    else:
        return "No"    
    
    # if naics in top_naics:
    #     return "Yes"
    # else:
    #     return "No"

def check_strong_naics(df, contract_no) -> str:
    """
    Check if the NAICS code value is one of the Top 30% of unique NAICS identified by the amount of SB Actions or SB Dollars.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str:  If NAICS is present, return "Yes". Otherwise, return "No".
    """
    
    cleansed_file_df = pd.read_csv(common_folders['cleansed_data_source_file'], usecols=['NAICS', 'Size Status', 'SB Dollars'])
    
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
    
    # Idnentify the number of uniqe values in the 'NAICS' column by count !!MAY HAVE TO ALSO KEEP IT THE LAST 3 YEARS OF DATA!!
    naics_count = cleansed_file_df['NAICS'].nunique()
    # print("Number of unique NAICS values:", naics_count)
    # print("")
    
    # Get the count for top 30% of NAICS based on SB Dollars or SB actions
    strong_naics_count = int(naics_count * .3)
    # print("Top 30% of NAICS count:", strong_naics_count)
    # print("")
    
    # Make sure the 'SB Dollars' column is numeric from the DataFrame argument (NOT THE FULL RAW CLEANSED FILE).  This value will be used to compare against the percentiles from the cleansed_file_df.
    # Ensure 'SB Dollars' column is string before removing currency formatting
    cleansed_file_df['SB Dollars'] = cleansed_file_df['SB Dollars'].astype(str)

    # Remove currency formatting from the SB Dollars column
    cleansed_file_df['SB Dollars'] = cleansed_file_df['SB Dollars'].str.replace('$', '').str.replace(',', '')

    # Convert the SB Dollars column to numeric
    cleansed_file_df['SB Dollars'] = pd.to_numeric(cleansed_file_df['SB Dollars'], errors='coerce')
    
    # Ensure the 'NAICS' column is of the same type as naics and only six digits are used
    cleansed_file_df['NAICS'] = cleansed_file_df['NAICS'].astype(str).str[:6].astype(type(naics))
    
    # Filter the cleased_file_df based on the naics identified from the contract being processed and print the filtered result to check if it matches the expected rows
    # filtered_df = cleansed_file_df.loc[(cleansed_file_df['NAICS'] == naics) & (cleansed_file_df['Size Status'] == "SB")]
    filtered_df = cleansed_file_df.loc[(cleansed_file_df['Size Status'] == "SB")]
    # print("Filtered DataFrame based on NAICS and Size Status:", filtered_df)
    # print("")
    
    # Define the targeted NAICS values based on the top 30% of unique NAICS based on total SB Dollars
    filtered_df = filtered_df.groupby('NAICS').agg({'SB Dollars': 'sum'}).reset_index()
    filtered_df = filtered_df.sort_values(by=['SB Dollars'], ascending=False)
    strong_naics_by_dollars = filtered_df['NAICS'].head(strong_naics_count).tolist()
    # print("Top NAICS by SB Dollars:", strong_naics_by_dollars)
    # print("")
    
    # Define the targeted NAICS values based on the top 30% of unique NAICS based on total SB Actions (determined by "Size Status" column).
    filtered_df = cleansed_file_df.loc[(cleansed_file_df['Size Status'] == "SB")]
    filtered_df = filtered_df.groupby('NAICS').agg({'Size Status': 'count'}).reset_index()
    filtered_df = filtered_df.sort_values(by=['Size Status'], ascending=False)
    strong_naics_by_actions = filtered_df['NAICS'].head(strong_naics_count).tolist()
    # print("Top NAICS by SB Actions:", strong_naics_by_actions)
    # print("")
    
    # top_naics = ['541611', '541512', '541330']
    
    # Check if the 'NAICS' value is in either top_naics_by_actions or top_naics_by_actions. If yes in either one, return "Yes". If not in either, return "No"
    if naics in strong_naics_by_dollars or naics in strong_naics_by_actions:
        return "Yes"
    else:
        return "No"    
    
    # if naics in top_naics:
    #     return "Yes"
    # else:
    #     return "No"

# def check_weak_naics(df, contract_no) -> str:
    """
    Check if the NAICS code value is one of the Weak NAICS in ACC-RI by 10th percentile of SB Dollars or SB Actions.
    
    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.
    
    returns:
    str:  If NAICS is present in the 10th percentile, return "Yes". Otherwise, return "No".
    """
    # Read into df the data source file that will be used to develop the percentiles from the  Dollars column
    df = pd.read_csv(common_folders['cleansed_data_source_file'], usecols=['Contract No', 'NAICS', 'Size Status'])
        
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
    
    # Filter the DataFrame for rows where 'Size Status' is "SB"
    filtered_df = df[df['Size Status'] == "SB"]

    # Group by 'NAICS' and count the occurrences
    naics_count_df = filtered_df.groupby('NAICS').size().reset_index(name='count')
    
    # Calculate the 10th percentile of the 'count' column
    tenth_percentile = np.percentile(naics_count_df['count'], 10)
    
    # Filter the 'NAICS' values that fall within the 10th percentile
    weak_naics = naics_count_df[naics_count_df['count'] <= tenth_percentile]['NAICS'].values
    # print("Weak NAICS:", weak_naics)
    # print("")
    
    if naics in weak_naics:
        return "Yes"
    else:
        return "No"

def check_weak_naics(df, contract_no) -> str:
    """
    Check if the NAICS code value is one of the Top 30% of unique NAICS identified by the amount of SB Actions or SB Dollars.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str:  If NAICS is present, return "Yes". Otherwise, return "No".
    """
    
    cleansed_file_df = pd.read_csv(common_folders['cleansed_data_source_file'], usecols=['NAICS', 'Size Status', 'SB Dollars'])
    
    # Select the 'NAICS' value from the DataFrame based on the current contract number being processed
    naics = df.loc[df['Contract No'] == contract_no, 'NAICS'].values[0]
    
    # makre sure naics is a string and only the first six digits are used
    naics = str(naics)[:6]
    
    # Idnentify the number of uniqe values in the 'NAICS' column by count !!MAY HAVE TO ALSO KEEP IT THE LAST 3 YEARS OF DATA!!
    naics_count = cleansed_file_df['NAICS'].nunique()
    # print("Number of unique NAICS values:", naics_count)
    # print("")
    
    # Get the count for top 30% of NAICS based on SB Dollars or SB actions
    weak_naics_count = int(naics_count * .3)
    # print("Top 30% of NAICS count:", strong_naics_count)
    # print("")
    
    # Make sure the 'SB Dollars' column is numeric from the DataFrame argument (NOT THE FULL RAW CLEANSED FILE).  This value will be used to compare against the percentiles from the cleansed_file_df.
    # Ensure 'SB Dollars' column is string before removing currency formatting
    cleansed_file_df['SB Dollars'] = cleansed_file_df['SB Dollars'].astype(str)

    # Remove currency formatting from the SB Dollars column
    cleansed_file_df['SB Dollars'] = cleansed_file_df['SB Dollars'].str.replace('$', '').str.replace(',', '')

    # Convert the SB Dollars column to numeric
    cleansed_file_df['SB Dollars'] = pd.to_numeric(cleansed_file_df['SB Dollars'], errors='coerce')
    
    # Ensure the 'NAICS' column is of the same type as naics and only six digits are used
    cleansed_file_df['NAICS'] = cleansed_file_df['NAICS'].astype(str).str[:6].astype(type(naics))
    
    # Filter the cleased_file_df based on the naics identified from the contract being processed and print the filtered result to check if it matches the expected rows
    # filtered_df = cleansed_file_df.loc[(cleansed_file_df['NAICS'] == naics) & (cleansed_file_df['Size Status'] == "SB")]
    filtered_df = cleansed_file_df.loc[(cleansed_file_df['Size Status'] == "SB")]
    # print("Filtered DataFrame based on NAICS and Size Status:", filtered_df)
    # print("")
    
    # Define the targeted NAICS values based on the top 30% of unique NAICS based on total SB Dollars
    filtered_df = filtered_df.groupby('NAICS').agg({'SB Dollars': 'sum'}).reset_index()
    filtered_df = filtered_df.sort_values(by=['SB Dollars'], ascending=True)
    weak_naics_by_dollars = filtered_df['NAICS'].head(weak_naics_count).tolist()
    # print("Top NAICS by SB Dollars:", strong_naics_by_dollars)
    # print("")
    
    # Define the targeted NAICS values based on the top 30% of unique NAICS based on total SB Actions (determined by "Size Status" column).
    filtered_df = cleansed_file_df.loc[(cleansed_file_df['Size Status'] == "SB")]
    filtered_df = filtered_df.groupby('NAICS').agg({'Size Status': 'count'}).reset_index()
    filtered_df = filtered_df.sort_values(by=['Size Status'], ascending=True)
    weak_naics_by_actions = filtered_df['NAICS'].head(weak_naics_count).tolist()
    # print("Top NAICS by SB Actions:", strong_naics_by_actions)
    # print("")
    
    # top_naics = ['541611', '541512', '541330']
    
    # Check if the 'NAICS' value is in either top_naics_by_actions or top_naics_by_actions. If yes in either one, return "Yes". If not in either, return "No"
    if naics in weak_naics_by_dollars or naics in weak_naics_by_actions:
        return "Yes"
    else:
        return "No"    
    
    # if naics in top_naics:
    #     return "Yes"
    # else:
    #     return "No"
       
# def check_forecast(df, contract_no) -> str:
    """
    Check if the contract is a forecasted action.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: "Yes" if the contract is a forecasted action, "No" otherwise.
    """
    forecast = df.loc[df['Contract No'] == contract_no, 'Forecast'].values[0]
    
    if forecast == 1:
        return "Yes"
    else:
        return "No"

def check_modification(df,contract_no) -> str:
    '''
    Check if the contract has a modification and get the most recent number identified by the award date.
    
    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.
    
    Returns:
    str: The most recent modification number.
    '''
    #Define the modifications file which will be the cleansed data source file
    modification_df = pd.read_csv(common_folders['cleansed_data_source_file'], usecols=['Contract No', 'Modification No', 'Award Date', 'Contract Action Type'])
    
    # Select the 'Contract No' value from the DataFrame based on the current contract number being processed
    contract_no = df.loc[df['Contract No'] == contract_no, 'Contract No'].values[0]
    
    # Filter the modifications_df based on the contract_no identified from the contract being processed where Contract Action Type is "MODIFICATION"
    modification_df = modification_df.loc[(modification_df['Contract No'] == contract_no) & (modification_df['Contract Action Type'] == "MODIFICATION")]
    
    # Sort the modifications_df by 'Award Date' in descending order
    modification_df = modification_df.sort_values(by='Award Date', ascending=False)
    
    # Get the most recent 'Modification Number' from the modifications_df
    if modification_df.empty:
        return "No Modifications"
    else:
        return modification_df['Modification No'].values[0]
    
def check_forecast(df, contract_no) -> str:
        '''
        Check if the contract is a forecasted action and return the VCE-PCF Cabinet Name.
        
        Args:
        df (pd.DataFrame): The DataFrame containing the data to be processed.
        contract_no (str): The contract number being processed.
        
        Returns:
        str: The VCE-PCF Cabinet Name.
        '''
        # Define the forecast file which will be the cleansed data source file
        forecast_df = pd.read_csv(common_folders['forecast_file'], usecols=['VCE-PCF Cabinet Name', 'FOLLOWON CONTRACT'])
        
        # Search the "FOLLOWON CONTRACT" column in the forecast_df for the contract_no identified from the contract being processed
        forecast_df = forecast_df.loc[forecast_df['FOLLOWON CONTRACT'] == contract_no]
        if contract_no in forecast_df['FOLLOWON CONTRACT'].values:
            return str(forecast_df['VCE-PCF Cabinet Name'].values[0])
        else:
            return "No Forecast Identified"
        
def check_pcf_cabinet_link(df, contract_no) -> str:
    '''
    Check the file and determine if there is a link to the identified contract being processes.
    
    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.
    
    Returns:
    str: The hyperlink to the PCF Cabinet.
    '''
    # Define the file where hyperlinks are stored
    hyperlink_df = pd.read_csv(common_folders['hyperlinks_file'], usecols= ['PCF Access', 'Contract', 'Order'])
    
    #Define Order No from argument df
    order_no = df.loc[df['Contract No'] == contract_no, 'Order No'].values[0] 
        
    # Define the base URL for the PCF Cabinet
    base_url = "https://pcf.army.mil/pcf/login.htm?t=1460636157811&cabinetName="

    # Construct the full URL by appending the contract number to the base URL
    # hyperlink = f"{base_url}{contract_no}"

    # Check if order no is in the hyperlink_df['Order'] column. If a value exists, return the value from the PCF Access column from the hyperlink_df.  If no, check if contract_no is in the hyperlink_df['Contract"] column. If a value exists, return the value of PCF Access column from hyperlink_df.  Otherwise, return naics not found message.
    if order_no in hyperlink_df['Order'].values:
        return hyperlink_df.loc[hyperlink_df['Order'] == order_no, 'PCF Access'].values[0]
        # return f"{base_url}{order_no}"
    elif contract_no in hyperlink_df['Contract'].values:
        return hyperlink_df.loc[hyperlink_df['Contract'] == contract_no, 'PCF Access'].values[0]
        # return f"{base_url}{contract_no}"
    else:
        return f'No PCF cabinet link found'
    
  
    return hyperlink

def check_it_buy(df, contract_no) -> str:
    """
    Check the NAICS Description, PSC Description, OMB Level 1 and OMB Level 2 columns for certain combinations and keywords to determine if it is an IT buy.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: "Yes" if the NAICS description contains specific keywords, "No" otherwise.
    """
    # Define the IT Buy keywords
    it_buy_keywords = ['IT', 'INFORMATION TECHNOLOGY', 'TECHNOLOGY', 'SOFTWARE', 'HARDWARE', 'COMPUTER', 'NETWORK', 'CYBERSECURITY', 'CLOUD', 'DATA', 'ANALYTICS', 'AI', 'ARTIFICIAL INTELLIGENCE', 'MACHINE LEARNING', 'ML', 'IOT', 'INTERNET OF THINGS', 'BLOCKCHAIN', 'CRYPTO', 'CRYPTOCURRENCY', 'DIGITAL', 'ELECTRONIC', 'TELECOMMUNICATIONS', 'TELECOMM', 'TELECOM', 'TELEPHONE', 'TELEPHONY', 'TELEPHONIC', 'TELEPHONICS']
    
    # Use columns 'NAICS', 'NAICS Description', 'PSC Decription', 'OMB Level 1', and 'OMB Level 2' to check for IT Buy keywords from df argument
    naics_description = df.loc[df['Contract No'] == contract_no, 'NAICS Description'].values[0]
    psc_description = df.loc[df['Contract No'] == contract_no, 'PSC Description'].values[0]
    omb_level_1 = df.loc[df['Contract No'] == contract_no, 'OMB Level 1'].values[0]
    omb_level_2 = df.loc[df['Contract No'] == contract_no, 'OMB Level 2'].values[0]
    
    # # Check if any of the IT Buy keywords are present in the 'NAICS Description', 'PSC Description', 'OMB Level 1', or 'OMB Level 2' columns
    # if any(keyword in naics_description.upper() for keyword in it_buy_keywords) or any(keyword in psc_description.upper() for keyword in it_buy_keywords) or any(keyword in omb_level_1.upper() for keyword in it_buy_keywords) or any(keyword in omb_level_2.upper() for keyword in it_buy_keywords):
    #     return "Yes"
    # else:
    #     return "No"
    
    # Create a regex pattern to match whole words only
    pattern = r'\b(?:' + '|'.join(re.escape(keyword) for keyword in it_buy_keywords) + r')\b'
    
    # Check if any of the IT Buy keywords are present in the 'NAICS Description', 'PSC Description', 'OMB Level 1', or 'OMB Level 2' columns
    if (re.search(pattern, naics_description, re.IGNORECASE) or
        re.search(pattern, psc_description, re.IGNORECASE) or
        re.search(pattern, omb_level_1, re.IGNORECASE) or
        re.search(pattern, omb_level_2, re.IGNORECASE)):
        return "Yes"
    else:
        return "No"

def check_socio_sole_source_eligible(df, contract_no) -> str:
    """
    Check if the contract is eligible for sole source award based on the dollar value.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    contract_no (str): The contract number being processed.

    Returns:
    str: "Yes" if the contract is eligible for sole source award based on the $4M threshold for SDVOSB and $4.5M for all others, "No" otherwise.
    """
    # Select the 'Socio-Economic Status' value from the DataFrame based on the current contract number being processed
    sdvosb_sole_source_threshold = 4000000
    all_others_sole_source_threshold = 4500000
    
    # create a list of all other sole source thresholds
    all_eligible_sole_source_categories = ['8(a)', 'HUBZone', 'WOSB', 'EDWOSB', 'SDVOSB']
    eligible_sole_sole_categories_above_sdvosb_threshold = ['8(a)', 'HUBZone', 'WOSB', 'EDWOSB']
    eligible_sole_source_categories_above_sdvosb_not_wosb_naics_list = ['8(a)', 'HUBZone']
    eligible_sole_sole_categories_not_on_wosb_naics_list = ['8(a)', 'HUBZone', 'SDVOSB']
    
    #Get dollar value from the 'SB Dollars' from df argument based on the contract_no being processed
    sb_dollars = df.loc[df['Contract No'] == contract_no, 'SB Dollars'].values[0]
    #Remove currency formatting from sb_dollars
    sb_dollars = sb_dollars.replace('$', '').replace(',', '')
    
    # Convert the 'SB Dollars' value to numeric
    sb_dollars = pd.to_numeric(sb_dollars, errors='coerce')    
        
    # # If the SB dollars less than sdvosb_sole_source_threshold AND the WOSB NAICS list is "Yes", return all_eligible_sole_source_categories as str
    # if sb_dollars > all_others_sole_source_threshold:
    #     return 'No'
    # elif sb_dollars >= sdvosb_sole_source_threshold and check_wosb_naics(df, contract_no) == "No":
    #     return ', '.join(eligible_sole_source_categories_above_sdvosb_not_wosb_naics_list)
    # elif sb_dollars <= sdvosb_sole_source_threshold and check_wosb_naics(df, contract_no) == "No":
    #     return ', '.join(eligible_sole_sole_categories_not_on_wosb_naics_list)
    # elif sb_dollars >= sdvosb_sole_source_threshold and check_wosb_naics(df, contract_no) != "No":
    #     return ', '.join(eligible_sole_sole_categories_above_sdvosb_threshold)
    # elif sb_dollars <= sdvosb_sole_source_threshold and check_wosb_naics(df, contract_no) != "No":
    #     return ', '.join(all_eligible_sole_source_categories)

sb_profile_analysis_functions = {
    "IT Buy" : check_it_buy, # Check NAICS desription to determine if it is an IT buy, search for specific keywords and return yes or no
    # "Strong Competition" : check_strong_competition, # Get a sense of average number of offerors against this NAICS (use all army data source file)
    "Size Standard" : check_size_standard,
    "Top NAICS" : check_top_naics, #Top 25 NIACS either by SB Dollars or SB Actions
    "Target NAICS" : check_targeted_naics, #NAICS identified by specific needs or objectives or rationales/logic
    "WOSB Eligible" : check_wosb_naics, # Check if the NAICS code value is present in the Underrepresented WOSB NAICS listing
    "Strong NAICS" : check_strong_naics, #Top 30% of NAICS based on SB Dollars or SB Actions
    "Weak NAICS" : check_weak_naics, #10th percentile of SB Dollars or SB Actions
    # "Socio SS Eligible" : check_socio_sole_source_eligible, # Check if the $ value is below the threshold ($4M SDVOSB, $4.5M all others)
    "ACC RI Awards" : check_acc_ri_awards, #Awards that went to SB under the identified NAICS
    "All ACC Awards" : check_all_acc_awards, #All awards made by ACC across the enterprise
    "Awardee SB" : check_if_awardee_sb,
    "Awardee Socio" : check_awardee_socioeconomic_status,
    "NMR Waiver Available" : check_if_nmr_waiver_available, #Does an NMR waiver exist based on NAICS
    "Financial Risk" : check_financial_risk, #Financial risk to industry based on distribution of SB awards under identified NAICS"
    "Modification No" : check_modification, #Check if the contract is a modification and get the most recent number
    "PCF Cabinet" : check_pcf_cabinet_link, #Provide link to PCF Cabinet and return a str or hyperlink
    "Forecast No" : check_forecast # Identify the forecast solicitation/PANCOC number
}   
                          
def create_contract_profiles():
    """
    Populate the Contract Details tables based on the Contract Profile Data Elements dictionary.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    template_path (str): The path to the template document.
    completed_folder (str): The path to the folder where completed documents will be saved.

    Returns:
    None
    """
    # Define the file to needed to create the dataframe (df)
    insights_target1_file = find_latest_file(common_folders['insights_target1_folder'], '*.csv')
    
    # Select the insights_target1.csv file from the insights_target1_folder
    df = pd.read_csv(insights_target1_file)
     
    # Define the completed folder location
    completed_profiles = common_folders["completed_profiles_folder"]
    
    # Define the location of the template document
    template_file = os.path.join(common_folders["contract_profiles_folder"], 'template_contract_profile.docx')

    # Ensure the completed_profiles folder exists
    if not os.path.exists(completed_profiles):
        os.makedirs(completed_profiles)

    # Iterate over each row in the DataFrame
    # Set the maximum number of rows to process for testing
    max_rows = 10  
    row_count = 0

    for index, row in df.iterrows():
        # Activate the following for testing see max row above
        if row_count >= max_rows:
            break

        # Load the template document
        template = Document(template_file)

        # Create a copy of the template. !!!!!Eventually remove "Target row and date" from the filename and simply replace the file when the script is rerun.  This will ensure that every contract profile is updated quarterly with any new data that can be autopopulated.  Make sure it only updates the tables and not the entire document.
        contract_no = row["Contract No"].replace('/', '_')
        today_date = datetime.datetime.now().strftime('%Y-%m-%d')
        new_filename = f'Target_{index+1}_{contract_no}_{today_date}.docx'
        new_filepath = os.path.join(completed_profiles, new_filename)
        template.save(new_filepath)

        # Load the copied template document and start populating the contract details
        doc = Document(new_filepath)

        # Create a uniform table for the Contract Details information in the document with four columns
        table1 = doc.add_table(rows=1, cols=4)
        table1.style = 'Table Grid'
        
        # Remove the first (and only) row which was added by default.  This will be replaced by the header row.  This may break if libraries are updated so be aware.    
        if len(table1.rows) > 1:
            table1._tbl.remove(table1.rows[0]._tr)
        
        # Add a header row and merge cells and center align for the "Contract Details" header
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1]).merge(hdr_cells[2]).merge(hdr_cells[3])
        hdr_cells[0].text = 'Contract Details'
        hdr_cells[0].paragraphs[0].alignment = 1  # Center alignment
        set_cell_font(hdr_cells[0], font_size=11, bold=True)
        
        # Populate the table with all values from contract_profile_data_elements based on the table1 structure defined above
        values = list(contract_profile_data_elements.values())
        for i in range(0, len(values), 2):
            row_cells = table1.add_row().cells
            row_cells[0].text = values[i]
            if values[i] in sb_profile_analysis_functions:
                row_cells[1].text = sb_profile_analysis_functions[values[i]](df, contract_no)
            else:
                row_cells[1].text = str(row[values[i]]) if values[i] in df.columns else 'Data Gap'
            if i + 1 < len(values):
                row_cells[2].text = values[i + 1]
                if values[i + 1] in sb_profile_analysis_functions:
                    row_cells[3].text = sb_profile_analysis_functions[values[i + 1]](df, contract_no)
                else:
                    row_cells[3].text = str(row[values[i + 1]]) if values[i + 1] in df.columns else 'Data Gap'
            else:
                row_cells[2].text = ''
                row_cells[3].text = ''
            for cell in row_cells:
                set_cell_font(cell)
        
        # # Add a row titled "PCF Cabinet" in the first cell and the remaining 3 cells are merge and left blank
        # row_cells = table1.add_row().cells
        # row_cells[0].text = 'PCF Cabinet'
        # row_cells[1].merge(row_cells[2]).merge(row_cells[3])
        # row_cells[1].text = sb_profile_analysis_functions['PCF Cabinet'](df, contract_no)
        # for cell in row_cells:
        #     set_cell_font(cell)
        
        # # Add a row titled "Forecast No" in the first cell and the remaining 3 cells are merge and left blank    
        # row_cells = table1.add_row().cells
        # row_cells[0].text = 'Forecast No'
        # row_cells[1].merge(row_cells[2]).merge(row_cells[3])
        # row_cells[1].text = sb_profile_analysis_functions['Forecast No'](df, contract_no)
        # for cell in row_cells:
        #    set_cell_font(cell)
                         
        # Add a paragraph break to ensure the next table is not connected to the previous one
        doc.add_paragraph()
        
        # Create a uniform table for the Small Business Profile Analysis information in the document with four columns
        table2 = doc.add_table(rows=1, cols=6)
        table2.style = 'Table Grid'

        # Add a header row and merge cells and center align for the "Small Business Profile Analysis" header
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1]).merge(hdr_cells[2]).merge(hdr_cells[3].merge(hdr_cells[4]).merge(hdr_cells[5]))
        hdr_cells[0].text = 'Small Business Profile Analysis'
        hdr_cells[0].paragraphs[0].alignment = 1  # Center alignment
        set_cell_font(hdr_cells[0], font_size=11, bold=True)

        # Populate the table with all values from contract_profile_data_elements based on the table1 structure defined above.  Also, check to see if the specific value has a function to populate the cell.
        values = list(sb_profile_analysis_elements.values())  
        for i in range(0, len(values), 3): 
            row_cells = table2.add_row().cells
            row_cells[0].text = values[i]
            # Check if the value has a function associated with it to populate the cell
            if values[i] in sb_profile_analysis_functions:
                row_cells[1].text = sb_profile_analysis_functions[values[i]](df, contract_no)
            else:
                row_cells[1].text = str(row[values[i]]) if values[i] in df.columns else ''
            # Check if the value + 1 has a function associated with it to populate the cell
            if i + 1 < len(values):
                row_cells[2].text = values[i + 1]
                if values[i + 1] in sb_profile_analysis_functions:
                    row_cells[3].text = sb_profile_analysis_functions[values[i + 1]](df, contract_no)
                else:
                    row_cells[3].text = str(row[values[i + 1]]) if values[i + 1] in df.columns else ''
            else:
                row_cells[2].text = ''
                row_cells[3].text = ''
            # Check if the value +2 has a function associated with it to populate the cell
            if i + 2 < len(values):
                row_cells[4].text = values[i + 2]
                if values[i + 2] in sb_profile_analysis_functions:
                    row_cells[5].text = sb_profile_analysis_functions[values[i + 2]](df, contract_no)
                else:
                    row_cells[5].text = str(row[values[i + 2]]) if values[i + 2] in df.columns else ''
            else:
                row_cells[4].text = ''
                row_cells[5].text = ''
            for cell in row_cells:
                set_cell_font(cell)
        
        # Add a final row titled "Remarks" in the first cell and the remaining 3 cells are mergec and left blank
        row_cells = table2.add_row().cells
        row_cells[0].text = 'Remarks'
        row_cells[1].merge(row_cells[2]).merge(row_cells[3].merge(row_cells[4]).merge(row_cells[5]))
        for cell in row_cells:
            set_cell_font(cell)
        
        # Add a paragraph break to ensure the sections is not connected to the previous one
        doc.add_paragraph()
        
        # Save the populated document
        doc.save(new_filepath)
        print(f"Populated contract details saved to: {new_filepath}")

        row_count += 1
    # Create a log of completed contracts and check the amount of documents compared to the amount of rows in the DataFrame
    log_file = os.path.join(completed_profiles, 'completed_profiles_log.txt')
    with open(log_file, 'a') as log:
        log.write(f"Completed profiles on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\nCompleted contracts location: {completed_profiles}\n\n")  

# def update_contract_profiles_tables():
    # TBD

def insights_test2(df):
    """Make a copy of the cleaned data and move it to the "insight2" folder.

    Args:
        df (_type_): _description_
    """
    # Define the insight2 folder location.
    insight2_folder = r'C:/GitHub/contract_profiles/data/processed/insight2'
    
    # If the insight2 folder does not exist, create it.
    if not os.path.exists(insight2_folder):
        os.makedirs(insight2_folder)
    
    # Create a copy of the df file and save it to the insight2 folder.
    df.to_csv(os.path.join(insight2_folder, 'insight2.csv'), index=False)
    print(f"Data copied to {insight2_folder}")
    
    return df

def check_format_rows():
    """Check the format of the rows in the DataFrame.

    Args:
        df (_type_): _description_

    Returns:
        _type_: _description_
    """
    
     # Select the insight_target.csv file from the targets folder
    df = pd.read_csv(common_folders["insights_target1_folder", 'insights_target1.csv'])
    
    # Identify the format of each column in the dfand connect it to the contract data elements dictionary and create a list of the format of the rows for each contract data element value
    format_list = []
    for value in contract_profile_data_elements.values():
        if value in df.columns:
            format_list.append(f"{value}: {df[value].dtype}")
            
    print(format_list)