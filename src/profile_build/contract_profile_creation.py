from dconfig import common_folders, pd, os, datetime, find_latest_file, contract_profile_data_elements, sb_profile_analysis_functions, sb_profile_analysis_elements, set_cell_font
from docx import Document
from utils.gen_utils import check_if_folder_exist, does_file_exist

def load_inisght_data(insight_file) -> pd.DataFrame:
    """Loads the data based on the specifically insight argument."""
    
    insight_df = pd.read_csv(insight_file)
        
    return insight_df

def get_new_document_path(completed_profiles_path, contract_no, index):
    """Generate a path for a new contract profile document."""
    
    today_date = datetime.datetime.now().strftime('%Y-%m-%d')
    
    filename = f'{contract_no.replace("/", "_")}_{today_date}.docx'
    
    return os.path.join(completed_profiles_path, filename)

def copy_profile_template(template_file, formal_template) -> str:
    """Prepare a new document based on the template document."""
    # Load the template document
    template = Document(template_file)

    #Save a copy of the template document
    template.save(formal_template)

    return formal_template

def ensure_directory_exists(directory):
    """Ensure the specified directory exists."""
    
    if not os.path.exists(directory):
        os.makedirs(directory)

def create_contract_details_table(formal_template) -> None:
    """Create the Contract Details table."""
    
    table = formal_template.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    if len(table.rows) > 1:
        table._tbl.remove(table.rows[0]._tr)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[1]).merge(hdr_cells[2]).merge(hdr_cells[3])
    
    hdr_cells[0].text = 'Contract Details'
    hdr_cells[0].paragraphs[0].alignment = 1  # Center alignment
    
    set_cell_font(hdr_cells[0], font_size=11, bold=True)
    
    return table

def create_sb_profile_analysis_table(formal_template):
    """Create and return the Small Business Profile Analysis table."""
    table = formal_template.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[1]).merge(hdr_cells[2]).merge(hdr_cells[3].merge(hdr_cells[4]).merge(hdr_cells[5]))
    
    hdr_cells[0].text = 'Small Business Profile Analysis'
    hdr_cells[0].paragraphs[0].alignment = 1  # Center alignment
    
    set_cell_font(hdr_cells[0], font_size=11, bold=True)
    return table

def create_contract_profiles():
    """
    Populate the Contract Details tables based on the Contract Profile Data Elements dictionary.

    Args:
    df (pd.DataFrame): The DataFrame containing the data to be processed.
    template_path (str): The path to the template document.
    completed_folder (str): The path to the folder where completed documents will be saved.

    Returns:
    None
    """
    # Define the file to needed to create the dataframe (df)
    insights_target1_file = find_latest_file(common_folders['insights_target1_folder'], '*.csv')
    
    # Select the insights_target1.csv file from the insights_target1_folder
    df = pd.read_csv(insights_target1_file)
     
    # Define the completed folder location
    completed_profiles = common_folders["completed_profiles_folder"]
    
    # Define the location of the template document
    template_file = os.path.join(common_folders["contract_profiles_folder"], 'template_contract_profile.docx')

    # Ensure the completed_profiles folder exists
    if not os.path.exists(completed_profiles):
        os.makedirs(completed_profiles)

    # Iterate over each row in the DataFrame
    # Set the maximum number of rows to process for testing
    max_rows = 10  
    row_count = 0

    for index, row in df.iterrows():
        # Activate the following for testing see max row above
        if row_count >= max_rows:
            break

        # Load the template document
        template = Document(template_file)

        # Create a copy of the template. !!!!!Eventually remove "Target row and date" from the filename and simply replace the file when the script is rerun.  This will ensure that every contract profile is updated quarterly with any new data that can be autopopulated.  Make sure it only updates the tables and not the entire document.
        contract_no = row["Contract No"].replace('/', '_')
        today_date = datetime.datetime.now().strftime('%Y-%m-%d')
        new_filename = f'Target_{index+1}_{contract_no}_{today_date}.docx'
        new_filepath = os.path.join(completed_profiles, new_filename)
        template.save(new_filepath)

        # Load the copied template document and start populating the contract details
        doc = Document(new_filepath)

        # Create a uniform table for the Contract Details information in the document with four columns
        table1 = doc.add_table(rows=1, cols=4)
        table1.style = 'Table Grid'
        
        # Remove the first (and only) row which was added by default.  This will be replaced by the header row.  This may break if libraries are updated so be aware.    
        if len(table1.rows) > 1:
            table1._tbl.remove(table1.rows[0]._tr)
        
        # Add a header row and merge cells and center align for the "Contract Details" header
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1]).merge(hdr_cells[2]).merge(hdr_cells[3])
        hdr_cells[0].text = 'Contract Details'
        hdr_cells[0].paragraphs[0].alignment = 1  # Center alignment
        set_cell_font(hdr_cells[0], font_size=11, bold=True)
        
        # Populate the table with all values from contract_profile_data_elements based on the table1 structure defined above
        values = list(contract_profile_data_elements.values())
        for i in range(0, len(values), 2):
            row_cells = table1.add_row().cells
            row_cells[0].text = values[i]
            if values[i] in sb_profile_analysis_functions:
                row_cells[1].text = sb_profile_analysis_functions[values[i]](df, contract_no)
            else:
                row_cells[1].text = str(row[values[i]]) if values[i] in df.columns else 'Data Gap'
            if i + 1 < len(values):
                row_cells[2].text = values[i + 1]
                if values[i + 1] in sb_profile_analysis_functions:
                    row_cells[3].text = sb_profile_analysis_functions[values[i + 1]](df, contract_no)
                else:
                    row_cells[3].text = str(row[values[i + 1]]) if values[i + 1] in df.columns else 'Data Gap'
            else:
                row_cells[2].text = ''
                row_cells[3].text = ''
            for cell in row_cells:
                set_cell_font(cell)
        
        # # Add a row titled "PCF Cabinet" in the first cell and the remaining 3 cells are merge and left blank
        # row_cells = table1.add_row().cells
        # row_cells[0].text = 'PCF Cabinet'
        # row_cells[1].merge(row_cells[2]).merge(row_cells[3])
        # row_cells[1].text = sb_profile_analysis_functions['PCF Cabinet'](df, contract_no)
        # for cell in row_cells:
        #     set_cell_font(cell)
        
        # # Add a row titled "Forecast No" in the first cell and the remaining 3 cells are merge and left blank    
        # row_cells = table1.add_row().cells
        # row_cells[0].text = 'Forecast No'
        # row_cells[1].merge(row_cells[2]).merge(row_cells[3])
        # row_cells[1].text = sb_profile_analysis_functions['Forecast No'](df, contract_no)
        # for cell in row_cells:
        #    set_cell_font(cell)
                         
        # Add a paragraph break to ensure the next table is not connected to the previous one
        doc.add_paragraph()
        
        # Create a uniform table for the Small Business Profile Analysis information in the document with four columns
        table2 = doc.add_table(rows=1, cols=6)
        table2.style = 'Table Grid'

        # Add a header row and merge cells and center align for the "Small Business Profile Analysis" header
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1]).merge(hdr_cells[2]).merge(hdr_cells[3].merge(hdr_cells[4]).merge(hdr_cells[5]))
        hdr_cells[0].text = 'Small Business Profile Analysis'
        hdr_cells[0].paragraphs[0].alignment = 1  # Center alignment
        set_cell_font(hdr_cells[0], font_size=11, bold=True)

        # Populate the table with all values from contract_profile_data_elements based on the table1 structure defined above.  Also, check to see if the specific value has a function to populate the cell.
        values = list(sb_profile_analysis_elements.values())  
        for i in range(0, len(values), 3): 
            row_cells = table2.add_row().cells
            row_cells[0].text = values[i]
            # Check if the value has a function associated with it to populate the cell
            if values[i] in sb_profile_analysis_functions:
                row_cells[1].text = sb_profile_analysis_functions[values[i]](df, contract_no)
            else:
                row_cells[1].text = str(row[values[i]]) if values[i] in df.columns else ''
            # Check if the value + 1 has a function associated with it to populate the cell
            if i + 1 < len(values):
                row_cells[2].text = values[i + 1]
                if values[i + 1] in sb_profile_analysis_functions:
                    row_cells[3].text = sb_profile_analysis_functions[values[i + 1]](df, contract_no)
                else:
                    row_cells[3].text = str(row[values[i + 1]]) if values[i + 1] in df.columns else ''
            else:
                row_cells[2].text = ''
                row_cells[3].text = ''
            # Check if the value +2 has a function associated with it to populate the cell
            if i + 2 < len(values):
                row_cells[4].text = values[i + 2]
                if values[i + 2] in sb_profile_analysis_functions:
                    row_cells[5].text = sb_profile_analysis_functions[values[i + 2]](df, contract_no)
                else:
                    row_cells[5].text = str(row[values[i + 2]]) if values[i + 2] in df.columns else ''
            else:
                row_cells[4].text = ''
                row_cells[5].text = ''
            for cell in row_cells:
                set_cell_font(cell)
        
        # Add a final row titled "Remarks" in the first cell and the remaining 3 cells are mergec and left blank
        row_cells = table2.add_row().cells
        row_cells[0].text = 'Remarks'
        row_cells[1].merge(row_cells[2]).merge(row_cells[3].merge(row_cells[4]).merge(row_cells[5]))
        for cell in row_cells:
            set_cell_font(cell)
        
        # Add a paragraph break to ensure the sections is not connected to the previous one
        doc.add_paragraph()
        
        # Save the populated document
        doc.save(new_filepath)
        print(f"Populated contract details saved to: {new_filepath}")

        row_count += 1
    # Create a log of completed contracts and check the amount of documents compared to the amount of rows in the DataFrame
    log_file = os.path.join(completed_profiles, 'completed_profiles_log.txt')
    with open(log_file, 'a') as log:
        log.write(f"Completed profiles on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\nCompleted contracts location: {completed_profiles}\n\n")