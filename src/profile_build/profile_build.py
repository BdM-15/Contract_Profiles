from ..dconfig import pd, os, datetime, data_folders, get_file_path, contract_profile_data_elements, sb_profile_analysis_elements, common_settings
from ..sb_analysis import *
from docx import Document
from ..utils.gen_utils import does_folder_exist, does_file_exist, get_file_name, set_cell_font
import shutil

def get_insight_data(insight_file) -> pd.DataFrame:
    """Loads the data based on the specifically insight argument."""
    
    insight_df = pd.read_csv(insight_file)
        
    return insight_df

def get_document_path(completed_profiles_path, contract_no):
    """Generate a path for a new contract profile document based on the insight being processed."""
    
    today_date = datetime.datetime.now().strftime('%Y-%m-%d')
    
    filename = f'{contract_no.replace("/", "_")}_{today_date}.docx'
    
    return os.path.join(completed_profiles_path, filename)

def copy_profile_template(template_file, contract_profile) -> Document:
    """Prepare a new document based on the template document."""
    # Copy the template to the new location
    shutil.copy(template_file, contract_profile)

    return contract_profile

def create_table(profile, table_name, columns, merge_cells=None) -> Document:
    """
    Create the Contract Details table using the data elements from the insight_dictionary argument.
    
    Args:
    profile (docx.Document): The document to which the table will be added.
    table_name (str): The name of the table being created.
    columns (int): The number of columns in the table.
    merge_cells (list): A list of tuples containing the cell ranges to be merged.  Default is None. Use True as argument to merge all cells.
    
    Returns:
    Table: The table that will contain the Contract Details.
    """
    table = profile.add_table(rows=1, cols=columns)
    table.style = 'Table Grid'
    if len(table.rows) > 1:
        table._tbl.remove(table.rows[0]._tr)
    hdr_cells = table.rows[0].cells
    if merge_cells:
        for merge_range in merge_cells:
            hdr_cells[merge_range[0]].merge(hdr_cells[merge_range[1]])
    hdr_cells[0].text = table_name
    hdr_cells[0].paragraphs[0].alignment = 1  # Center alignment
    set_cell_font(hdr_cells[0], font_size=11, bold=True)
    return table

def populate_table(table, df_row, contract_no, elements, elements_per_row, functions, add_remarks=False) -> Document:
    """
    Populate the table with data.
    
    Args:
    table (Table): The table to populate.
    df (pd.DataFrame): The DataFrame containing the data to populate the table.
    contract_no (str): The contract number to use to populate the table.
    elements (dict): The elements to populate the table with.
    functions (dict): The functions to use to populate the table.
    add_remarks (bool): Whether to add a "Remarks" row to the table. Default is False.
    
    Returns:
    Table: The populated table.
    """
    values = list(elements.values())
    num_columns =len(table.columns)
    
    for i in range(0, len(values), elements_per_row):
        row_cells = table.add_row().cells
        
        #Populate cells base on available columns
        for j in range(elements_per_row):
            if i + j < len(values):
                row_cells[2 * j].text = values[i + j]
                row_cells[2 * j + 1].text = functions.get(values[i + j], lambda df_row, cn: str(df_row[values[i + j]]) 
                                                          if values[i + j] in df_row else '')(df_row, contract_no)
            
        # Clear any remaining cells if they exist but are not needed
        for k in range(2* elements_per_row, num_columns):
            row_cells[k].text = ''
            
        for cell in row_cells:
            set_cell_font(cell)
    
    # Add a "Remarks" row if specified
    if add_remarks:
        row_cells = table.add_row().cells
        row_cells[0].text = 'Remarks'
        row_cells[1].merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4]).merge(row_cells[5])
        for cell in row_cells:
            set_cell_font(cell)
               
    return table

def generate_profiles(insight_file, completed_profiles_folder):
    """Generate profiles based on the insights."""
    does_folder_exist(completed_profiles_folder)
    
    insight_name = get_file_name(insight_file)
    insight_profiles = os.path.join(completed_profiles_folder, insight_name)
    does_folder_exist(insight_profiles)
    
    df = get_insight_data(insight_file)
    
    template_file = get_file_path('contract_profiles', 'profile_template')

    for index, df_row in df.iterrows():
        if index >= common_settings['max_rows']:  # Limit the number of profiles generated for testing
            break
        profile_location = get_document_path(insight_profiles, df_row["Contract No"].replace('/', '_'))
        copy_profile_template(template_file, profile_location)
        profile = Document(profile_location)

        # Create and populate the Contract Details table
        contract_details_table = create_table(profile, 'Contract Details', 4, [(0, 3)])
        populate_table(contract_details_table, df_row, df_row["Contract No"], contract_profile_data_elements, 2, sb_profile_analysis_functions, False)
    
        profile.add_paragraph('')
    
        # Create and populate the PCF and Foreast Link table
        pcf_forecast_table = create_table(profile, 'PCF and Forecast Link', 4, [(0, 1)])
        populate_table(pcf_forecast_table, df_row, df_row["Contract No"], {'1' : 'PCF Link', '2' : 'Forecast Link'}, 1, sb_profile_analysis_functions, False)
        
        profile.add_paragraph('')
        
        # Create and populate the Small Business Profile Analysis table
        sb_profile_analysis_table = create_table(profile, 'Small Business Profile Analysis', 6, [(0, 5)])
        populate_table(sb_profile_analysis_table, df_row, df_row["Contract No"], sb_profile_analysis_elements, 3, sb_profile_analysis_functions, True)

        # profile.add_paragraph('')
        
        profile.save(profile_location)
        print(f"Populated contract details saved to: {profile_location}")

    log_file = os.path.join(completed_profiles_folder, 'completed_profiles_log.txt')
    with open(log_file, 'a') as log:
        log.write(f"Completed profiles on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\nCompleted contracts location: {insight_profiles}\n\n")